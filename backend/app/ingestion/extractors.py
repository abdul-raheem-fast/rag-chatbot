import hashlib
import fitz  # pymupdf
import pandas as pd
import trafilatura
import httpx
from pathlib import Path
from docx import Document as DocxDocument
from openpyxl import load_workbook
from app.core.logging import get_logger

logger = get_logger(__name__)


def compute_hash(content: str) -> str:
    return hashlib.sha256(content.encode("utf-8")).hexdigest()


def extract_pdf(file_path: str) -> list[dict]:
    """Extract text from PDF, one entry per page."""
    pages = []
    doc = fitz.open(file_path)
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text").strip()
        if text:
            pages.append({
                "text": text,
                "metadata": {"page_number": page_num},
            })
    doc.close()
    if not pages:
        logger.warning("PDF extraction returned no text", file_path=file_path)
    return pages


def extract_csv(file_path: str) -> list[dict]:
    """Each row becomes a text entry with column headers as context."""
    df = pd.read_csv(file_path)
    entries = []
    columns = list(df.columns)
    for idx, row in df.iterrows():
        parts = [f"{col}: {row[col]}" for col in columns if pd.notna(row[col])]
        text = " | ".join(parts)
        if text.strip():
            entries.append({
                "text": text,
                "metadata": {"row_number": int(idx) + 1},
            })
    return entries


def extract_txt(file_path: str) -> list[dict]:
    """Read a plain text file."""
    path = Path(file_path)
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        logger.warning("TXT file is empty", file_path=file_path)
        return []
    return [{"text": text, "metadata": {"source_file": path.name}}]


def extract_docx(file_path: str) -> list[dict]:
    """Extract text from a .docx file, preserving paragraph structure."""
    doc = DocxDocument(file_path)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    if not paragraphs:
        # Fall back to tables if no paragraphs
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

    if not paragraphs:
        logger.warning("DOCX extraction returned no text", file_path=file_path)
        return []

    full_text = "\n\n".join(paragraphs)
    return [{"text": full_text, "metadata": {"source_file": Path(file_path).name}}]


def extract_xlsx(file_path: str) -> list[dict]:
    """Extract text from .xlsx — each row becomes an entry with column headers as context.
    Handles multiple sheets."""
    entries = []
    try:
        sheets = pd.read_excel(file_path, sheet_name=None, engine="openpyxl")
    except Exception as e:
        raise ValueError(f"Failed to read Excel file: {e}")

    for sheet_name, df in sheets.items():
        columns = list(df.columns)
        for idx, row in df.iterrows():
            parts = [f"{col}: {row[col]}" for col in columns if pd.notna(row[col])]
            text = " | ".join(parts)
            if text.strip():
                entries.append({
                    "text": text,
                    "metadata": {
                        "sheet_name": sheet_name,
                        "row_number": int(idx) + 1,
                    },
                })

    if not entries:
        logger.warning("XLSX extraction returned no data", file_path=file_path)
    return entries


def extract_website(url: str) -> list[dict]:
    """Scrape a URL and extract clean text."""
    downloaded = trafilatura.fetch_url(url)
    if not downloaded:
        raise ValueError(f"Could not fetch URL: {url}")
    text = trafilatura.extract(downloaded, include_links=False, include_tables=True)
    if not text:
        raise ValueError(f"Could not extract text from: {url}")
    return [{"text": text, "metadata": {"source_url": url}}]


async def extract_gdoc(doc_id: str, credentials_json: str) -> list[dict]:
    """Fetch Google Doc content via export URL.
    For public/shared docs: uses unauthenticated export link.
    For private docs: uses service account credentials if provided."""
    import json

    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
    headers = {}

    creds_data = None
    try:
        creds_data = json.loads(credentials_json) if credentials_json else None
    except (json.JSONDecodeError, TypeError):
        pass

    if creds_data and creds_data.get("type") == "service_account":
        try:
            from google.oauth2.service_account import Credentials
            scopes = ["https://www.googleapis.com/auth/drive.readonly"]
            creds = Credentials.from_service_account_info(creds_data, scopes=scopes)
            creds.refresh(google.auth.transport.requests.Request())
            headers["Authorization"] = f"Bearer {creds.token}"
        except Exception as e:
            logger.warning("Service account auth failed, trying public access", error=str(e))

    async with httpx.AsyncClient(follow_redirects=True) as client:
        resp = await client.get(export_url, headers=headers)
        if resp.status_code != 200:
            raise ValueError(
                f"Failed to fetch Google Doc {doc_id}: HTTP {resp.status_code}. "
                "Ensure the doc is shared publicly or a valid service account is configured."
            )
        text = resp.text.strip()

    if not text:
        raise ValueError(f"Google Doc {doc_id} returned empty content")
    return [{"text": text, "metadata": {"gdoc_id": doc_id}}]


async def extract_notion(page_id: str, api_token: str) -> list[dict]:
    """Fetch Notion page blocks and concatenate text.
    Uses the Notion API to get block children."""
    headers = {
        "Authorization": f"Bearer {api_token}",
        "Notion-Version": "2022-06-28",
    }
    url = f"https://api.notion.com/v1/blocks/{page_id}/children?page_size=100"
    texts = []
    async with httpx.AsyncClient() as client:
        resp = await client.get(url, headers=headers)
        if resp.status_code != 200:
            raise ValueError(f"Notion API error: {resp.status_code}")
        data = resp.json()
        for block in data.get("results", []):
            block_type = block.get("type", "")
            block_data = block.get(block_type, {})
            rich_texts = block_data.get("rich_text", [])
            for rt in rich_texts:
                plain = rt.get("plain_text", "")
                if plain.strip():
                    texts.append(plain)

    full_text = "\n".join(texts)
    if not full_text.strip():
        raise ValueError(f"No text found in Notion page {page_id}")
    return [{"text": full_text, "metadata": {"notion_page_id": page_id}}]
